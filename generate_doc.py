from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_cyberverse_doc():
    doc = Document()
    
    # Title
    title = doc.add_heading('Cyber Verse: Cybersecurity Training Crossword\nTechnical & Logical Architecture Document', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('Prepared for: Project Supervisors / Stakeholders')
    doc.add_paragraph('Subject: Comprehensive overview of the game mechanics, real-time backend synchronization, and administrative workflows.\n')

    # 1. Executive Summary
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph('Cyber Verse is a real-time, browser-based, multiplayer cybersecurity crossword simulation. Designed with a cyberpunk-inspired interface, it challenges teams of agents to "decrypt" a network grid by answering cybersecurity-related questions. The game utilizes a serverless architecture powered by Firebase Realtime Firestore to instantly synchronize game state, scoring, and administrative controls across all active clients without page reloads.')

    # 2. Application Architecture
    doc.add_heading('2. Application Architecture', level=1)
    
    doc.add_heading('2.1 Frontend Interface', level=2)
    doc.add_paragraph('The frontend is built using standard web technologies (HTML5, Vanilla JavaScript, and Tailwind CSS) to ensure maximum compatibility and performance. It features a responsive grid engine capable of rendering dynamic crossword structures up to 20x30 tiles.')
    
    doc.add_heading('2.2 Backend & State Management', level=2)
    doc.add_paragraph('Firebase Firestore serves as the central data hub. The game utilizes the "onSnapshot" real-time listeners to maintain strict parity between the server and all connected clients. The database is divided into two primary collections:')
    ul1 = doc.add_paragraph(style='List Bullet')
    ul1.add_run('Teams Collection: ').bold = True
    ul1.add_run('Stores live data for each registered team, including their current score, time elapsed, solved nodes, hints used, and submission status.')
    ul2 = doc.add_paragraph(style='List Bullet')
    ul2.add_run('Settings Document: ').bold = True
    ul2.add_run('Maintains the global game state (e.g., gameStarted, gameEnded) and global administrative variables (hint penalties, point values, dynamic question overrides).')

    # 3. Player Workflow & Game Logic
    doc.add_heading('3. Player Workflow & Game Logic', level=1)
    
    doc.add_heading('3.1 Registration & Lobby Phase', level=2)
    doc.add_paragraph('Players initialize a session by registering a Team Name and the names of two participating agents. Upon registration, they are placed into the "Awaiting Overseer" state. In this phase, players can view the encrypted grid but cannot interact with it until the Administrator officially launches the mission.')

    doc.add_heading('3.2 Gameplay Mechanics', level=2)
    doc.add_paragraph('Once the game begins, the grid unlocks. Players navigate the grid by clicking on nodes to reveal the corresponding "Clue Prompt".')
    
    p = doc.add_paragraph()
    p.add_run('Scoring System:\n').bold = True
    p.add_run('- Base Points: Each correctly solved word awards a base of 100 points.\n')
    p.add_run('- Penalty Mechanism: Incorrect attempts penalize the team dynamically based on Overseer settings (default: -2 points per incorrect attempt).\n')
    
    p2 = doc.add_paragraph()
    p2.add_run('Hint System:\n').bold = True
    p2.add_run('If a team is stuck, they may request a hint. Hints provide an executable Python snippet that players must logically parse to discover the answer. Utilizing a hint incurs a flat penalty (default: -60 points) from the base score of that specific node.')

    doc.add_heading('3.3 Finish Conditions', level=2)
    doc.add_paragraph('When a team solves the entire grid, they are transitioned to a "Network Secured" standby screen, freezing their final time and score. They remain in this state until the Admin triggers the global End Game sequence.')

    # 4. Administrative Features (Overseer Console)
    doc.add_heading('4. Administrative Features (Overseer Console)', level=1)
    doc.add_paragraph('The Overseer Console provides the administrator with robust, real-time control over the entire event environment.')
    
    doc.add_heading('4.1 Session Management', level=2)
    p3 = doc.add_paragraph()
    p3.add_run('- Wipe Session (New Game): ').bold = True
    p3.add_run('The admin can obliterate all historical team data and reset the environment, effectively providing a clean slate for a new cohort of players.\n')
    p3.add_run('- Admin Lobby: ').bold = True
    p3.add_run('Prior to starting the game, the admin can monitor a live counter and list of all teams currently connecting to the server.\n')
    p3.add_run('- Launch Mission: ').bold = True
    p3.add_run('A single click globally unlocks the grid for all waiting agents simultaneously.')

    doc.add_heading('4.2 Environmental Controls', level=2)
    p4 = doc.add_paragraph()
    p4.add_run('- Dynamic Timers: ').bold = True
    p4.add_run('The admin can configure the game timer to either Count Up (Standard format) or Count Down (Pressure format) from a specified number of minutes (e.g., 60 minutes).\n')
    p4.add_run('- Rule Adjustments: ').bold = True
    p4.add_run('Point values, hint penalties, and attempt thresholds can be adjusted in real-time. Any changes instantly propagate to all active clients.')

    doc.add_heading('4.3 Content Overrides', level=2)
    doc.add_paragraph('Administrators are not restricted to hardcoded game files. Through the "Hint Settings" portal, the admin can modify the Python code snippets attached to any node dynamically. These overrides are saved to Firebase and instantly pushed to any player requesting a hint.')

    doc.add_heading('4.4 End Game & Slido-Style Podium', level=2)
    doc.add_paragraph('When the admin clicks "End Game," all player sessions are forcibly halted. Players are transitioned to an interactive podium view that displays:')
    p5 = doc.add_paragraph()
    p5.add_run('1. Their personal rank and final score.\n')
    p5.add_run('2. A visually celebrated Top 3 Podium.\n')
    p5.add_run('3. A full scrollable leaderboard of all runner-up teams.')
    
    doc.add_paragraph('Additionally, the administrator or organizers can instantly export the final global standings as a PDF directly from the Leaderboard panel for official record-keeping.')

    # Footer
    doc.add_paragraph('\n\n---')
    doc.add_paragraph('Document compiled autonomously for internal review and superior presentation.')
    
    doc.save('CyberVerse_Game_Logic.docx')
    print("Successfully generated CyberVerse_Game_Logic.docx")

if __name__ == '__main__':
    create_cyberverse_doc()
